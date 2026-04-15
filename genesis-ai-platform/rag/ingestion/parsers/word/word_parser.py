"""
Word 文档解析器

支持 .docx 格式，转换为 Markdown
"""

import logging
import mimetypes
import re
from io import BytesIO
from pathlib import Path
from typing import Tuple, Dict, Any, List, Callable, Optional

logger = logging.getLogger(__name__)


class WordParser:
    """
    Word 文档解析器
    
    特点：
    - 支持 .docx（OOXML 格式）
    - 转换为 Markdown 格式
    - 保留标题层级
    - 支持表格转换
    - 支持图片占位符
    
    不支持：
    - .doc（旧版二进制格式）
    """
    
    def __init__(self, with_image_placeholder: bool = False):
        """
        初始化 Word 解析器
        
        Args:
            with_image_placeholder: 是否生成图片占位符
        """
        self.with_image_placeholder = with_image_placeholder
    
    def is_available(self) -> bool:
        """检查 python-docx 是否可用"""
        try:
            import docx
            return True
        except ImportError:
            logger.warning("[WordParser] python-docx 未安装")
            return False
    
    def parse(self, file_buffer: bytes) -> Tuple[str, Dict[str, Any]]:
        """
        解析 Word 文档
        
        Args:
            file_buffer: Word 文件字节内容
        
        Returns:
            (markdown_text, metadata): Markdown 文本和元数据
        """
        logger.info("[WordParser] 开始解析 Word 文档")
        
        if self.with_image_placeholder:
            markdown_text, image_placeholders, embedded_images = self._parse_to_markdown(
                file_buffer,
                with_image_placeholder=True
            )
            
            metadata = {
                "parse_method": "word",
                "parser": "python-docx",
                "image_count": len(image_placeholders),
                "docx_image_placeholders": image_placeholders,  # 使用 docx_ 前缀
                "docx_embedded_images": embedded_images,  # 使用 docx_ 前缀
            }
        else:
            markdown_text = self._parse_to_markdown(
                file_buffer,
                with_image_placeholder=False
            )
            
            metadata = {
                "parse_method": "word",
                "parser": "python-docx",
            }
        
        logger.info(f"[WordParser] 解析完成，文本长度: {len(markdown_text)}")
        
        return markdown_text, metadata
    
    def _parse_to_markdown(
        self,
        file_buffer: bytes,
        with_image_placeholder: bool = False
    ):
        """
        将 Word 文档转换为 Markdown
        
        Args:
            file_buffer: Word 文件字节内容
            with_image_placeholder: 是否生成图片占位符
        
        Returns:
            如果 with_image_placeholder=True: (markdown, image_placeholders, embedded_images)
            否则: markdown
        """
        try:
            from docx import Document
        except ImportError:
            raise RuntimeError(
                "python-docx 未安装，无法解析 Word 文档。\n"
                "安装命令: pip install python-docx"
            )
        
        document = Document(BytesIO(file_buffer))
        paragraphs_iter = iter(document.paragraphs)
        tables_iter = iter(document.tables)
        lines: List[str] = []
        image_placeholders: List[Dict[str, str]] = []
        embedded_images: List[Dict[str, Any]] = []
        image_index = 0
        rid_to_image_id: Dict[str, str] = {}
        related_parts = getattr(document.part, "related_parts", {})
        
        def next_image_placeholder(rid: str) -> str:
            nonlocal image_index
            image_id = rid_to_image_id.get(rid)
            if not image_id:
                image_index += 1
                image_id = f"docx_img_{image_index:04d}"
                rid_to_image_id[rid] = image_id
                image_placeholders.append({
                    "id": image_id,
                    "rid": rid,
                    "status": "pending",
                    "source": "docx_embedded",
                })
            return f"![{image_id}](docx://embedded/{image_id})"
        
        def iter_block_items(parent):
            """按文档原始顺序遍历段落与表格"""
            for child in parent.element.body.iterchildren():
                if child.tag.endswith("}p"):
                    yield next(paragraphs_iter)
                elif child.tag.endswith("}tbl"):
                    yield next(tables_iter)
        
        # 遍历文档块
        for block in iter_block_items(document):
            if hasattr(block, "text"):  # Paragraph
                para_md = self._paragraph_to_markdown(
                    block,
                    with_image_placeholder=with_image_placeholder,
                    placeholder_builder=next_image_placeholder if with_image_placeholder else None,
                )
                if para_md:
                    lines.append(para_md)
            else:  # Table
                table_md = self._table_to_markdown(
                    block,
                    with_image_placeholder=with_image_placeholder,
                    placeholder_builder=next_image_placeholder if with_image_placeholder else None,
                )
                if table_md:
                    lines.append(table_md)
        
        markdown = "\n\n".join(lines).strip()
        
        if with_image_placeholder:
            # 提取图片数据
            for item in image_placeholders:
                rid = item["rid"]
                image_part = related_parts.get(rid)
                if not image_part:
                    continue
                
                blob = getattr(image_part, "blob", None)
                if not blob:
                    continue
                
                content_type = getattr(image_part, "content_type", "application/octet-stream")
                ext = mimetypes.guess_extension(content_type) or Path(str(getattr(image_part, "partname", ""))).suffix
                if not ext:
                    ext = ".bin"
                
                embedded_images.append({
                    "id": item["id"],
                    "rid": rid,
                    "content_type": content_type,
                    "ext": ext,
                    "size": len(blob),
                    "blob": blob,
                })
            
            return markdown, image_placeholders, embedded_images
        
        return markdown
    
    def _paragraph_to_markdown(
        self,
        paragraph,
        with_image_placeholder: bool = False,
        placeholder_builder: Optional[Callable] = None,
        is_in_table: bool = False,
    ) -> str:
        """段落转 Markdown"""
        text = (paragraph.text or "").strip()
        image_tokens = self._extract_paragraph_image_tokens(
            paragraph,
            with_image_placeholder,
            placeholder_builder
        )
        
        if not text and not image_tokens:
            return ""
        
        # 检测标题级别
        heading_level = self._extract_heading_level(paragraph)
        image_suffix = f" {' '.join(image_tokens)}" if image_tokens else ""
        
        if heading_level and not is_in_table:
            heading_text = text or "图片内容"
            return f"{'#' * heading_level} {heading_text}{image_suffix}"
        
        # 检测列表
        style_name = (getattr(paragraph.style, "name", "") or "").lower()
        if "list bullet" in style_name:
            list_text = text or "图片内容"
            return f"- {list_text}{image_suffix}"
        if "list number" in style_name:
            list_text = text or "图片内容"
            return f"1. {list_text}{image_suffix}"
        
        if text:
            return f"{text}{image_suffix}"
        return " ".join(image_tokens)
    
    def _extract_heading_level(self, paragraph) -> int:
        """提取标题级别（1-6）"""
        # 1. 第一优先级：从段落自身的直接属性获取 (Direct Formatting)
        try:
            p_pr = paragraph._p.pPr
            if p_pr is not None:
                from docx.oxml.ns import qn
                ol = p_pr.find(qn('w:outlineLvl'))
                if ol is not None and ol.get(qn('w:val')) is not None:
                    # 0 -> H1, 1 -> H2, ...
                    level = int(ol.get(qn('w:val'))) + 1
                    if 1 <= level <= 9:
                        return min(level, 6)
        except Exception:
            pass
            
        # 2. 第二优先级：从关联样式及其继承链中获取 (Style Inheritance)
        style = getattr(paragraph, "style", None)
        curr_style = style
        visited_styles = set() # 防止循环依赖（虽然 Word 理论上不允许）
        
        while curr_style and curr_style.style_id not in visited_styles:
            visited_styles.add(curr_style.style_id)
            try:
                from docx.oxml.ns import qn
                # 在样式的 pPr 中查找 outlineLvl
                style_pPr = curr_style.element.find(qn('w:pPr'))
                if style_pPr is not None:
                    ol = style_pPr.find(qn('w:outlineLvl'))
                    if ol is not None:
                        val = ol.get(qn('w:val'))
                        if val is not None:
                            level = int(val) + 1
                            if 1 <= level <= 9:
                                return min(level, 6)
            except Exception:
                pass
            
            # 向父级样式追溯
            curr_style = getattr(curr_style, "base_style", None)

        # 3. 第三优先级：基于名称匹配的启发式匹配 (Heuristic Naming)
        if not style:
            return 0
            
        style_name = (getattr(style, "name", "") or "").lower()
        style_id = (getattr(style, "style_id", "") or "").lower()
        
        def normalize(s: str) -> str:
            # 移除空格、下划线、中划线，方便统一匹配
            return re.sub(r"[\s_-]+", "", s)
            
        name_norm = normalize(style_name)
        id_norm = normalize(style_id)
        
        # 扩展的正则表达式，支持：
        # - Heading 1, Heading 2
        # - 标题1, 标题2
        # - 1级, 2级, 一级, 二级
        # - Level 1, Level 2
        # - 级别1, 级别2
        cn_num_map = {"一": 1, "二": 2, "三": 3, "四": 4, "五": 5, "六": 6, "七": 7, "八": 8, "九": 9}
        
        pattern = r"(heading|标题|level|级别|第)?([1-9]|一|二|三|四|五|六|七|八|九)(级)?"
        for target in [name_norm, id_norm]:
            match = re.search(pattern, target)
            if match:
                num_str = match.group(2)
                level = cn_num_map.get(num_str) or int(num_str)
                return min(level, 6)
            
        # 映射特定逻辑样式
        # Title/标题 -> H1
        title_names = {"title", "标题", "主题", "正文标题", "大标题"}
        if name_norm in title_names or id_norm in title_names:
            return 1
            
        # Subtitle/副标题 -> H2
        subtitle_names = {"subtitle", "副标题", "小标题"}
        if name_norm in subtitle_names or id_norm in subtitle_names:
            return 2
            
        return 0
            
        return 0
    
    def _table_to_markdown(
        self,
        table,
        with_image_placeholder: bool = False,
        placeholder_builder: Optional[Callable] = None,
    ) -> str:
        """表格转 Markdown"""
        rows: List[List[str]] = []
        
        for row in table.rows:
            cells = []
            for cell in row.cells:
                cell_paragraphs = []
                for paragraph in cell.paragraphs:
                    para_md = self._paragraph_to_markdown(
                        paragraph,
                        with_image_placeholder=with_image_placeholder,
                        placeholder_builder=placeholder_builder,
                        is_in_table=True,
                    )
                    if para_md:
                        cell_paragraphs.append(para_md)
                cell_text = self._escape_markdown_table_cell("<br>".join(cell_paragraphs))
                cells.append(cell_text)
            rows.append(cells)
        
        if not rows:
            return ""
        
        # 估算最大列数
        max_cols = max(len(row) for row in rows)
        
        # 补齐列数
        normalized_rows = [row + [""] * (max_cols - len(row)) for row in rows]
        
        # 过滤全空行（保留第一行）
        final_rows = [normalized_rows[0]]
        for row in normalized_rows[1:]:
            if any(cell.strip() for cell in row):
                final_rows.append(row)
        
        if not final_rows:
            return ""
        
        header = final_rows[0]
        data_rows = final_rows[1:]
        
        header_line = "| " + " | ".join(header) + " |"
        separator_line = "| " + " | ".join(["---"] * max_cols) + " |"
        body_lines = ["| " + " | ".join(row) + " |" for row in data_rows]
        
        return "\n".join([header_line, separator_line, *body_lines])
    
    @staticmethod
    def _escape_markdown_table_cell(value: str) -> str:
        """转义 Markdown 表格单元格"""
        return (value or "").replace("|", "\\|").strip()
    
    @staticmethod
    def _extract_paragraph_image_tokens(
        paragraph,
        with_image_placeholder: bool,
        placeholder_builder: Optional[Callable],
    ) -> List[str]:
        """提取段落中的图片占位符"""
        if not with_image_placeholder or placeholder_builder is None:
            return []
        
        rid_namespace = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}"
        tokens: List[str] = []
        
        # 提取图片
        blips = paragraph._element.xpath(".//*[local-name()='blip']")
        for blip in blips:
            rid = blip.get(f"{rid_namespace}embed")
            if rid:
                tokens.append(placeholder_builder(rid))
        
        # 提取旧版图片
        legacy_images = paragraph._element.xpath(".//*[local-name()='imagedata']")
        for image in legacy_images:
            rid = image.get(f"{rid_namespace}id")
            if rid:
                tokens.append(placeholder_builder(rid))
        
        return tokens
